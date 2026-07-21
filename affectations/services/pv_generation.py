import hashlib
import subprocess
from pathlib import Path

from django.conf import settings
from django.core.exceptions import ImproperlyConfigured
from django.utils import timezone

from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from affectations.models import PvAffectation
from affectations.services.pv_rules import build_pv_key


PV_TEMPLATE_NAME = "pv_affectation_template.docx"
ELECTRONIC_SIGNATURE_PREFIX = "Signé électroniquement par "


def format_date(value):
    return value.strftime("%d/%m/%Y") if value else ""


def format_decimal(value):
    return f"{value:,.2f}".replace(",", " ") if value is not None else ""


def build_reference_fonciere(dossier):
    parts = [
        dossier.trn,
        dossier.num_trn,
        dossier.indice_trn,
        dossier.num_id,
    ]
    return " / ".join(str(part) for part in parts if part)


def build_pv_context(dossier, administration):
    return {
        "num_dossier": dossier.num_dossier or "",
        "administration_nom_fr": administration.nom or "",
        "administration_nom_ar": administration.nom_ar or "",
        "adresse_fr": administration.adresse_fr or "",
        "adresse_ar": administration.adresse_ar or "",
        "denomination_projet": dossier.denomination_projet or "",
        "reference_fonciere": build_reference_fonciere(dossier),
        "superficie_concernee": format_decimal(dossier.superficie_concernee),
        "montant_affectation": format_decimal(dossier.montant_affectation),
        "numero_pv": dossier.numero_pv or "",
        "date_pv": format_date(dossier.date_envoi_pva_dr),
        "dr": dossier.dr or "",
        "delegation": dossier.delegation or "",
        "statut_pv": dossier.statut_pv or "",
        "type_pv": dossier.type_pv or "",
    }


def get_template_path(template_name=PV_TEMPLATE_NAME):
    return Path(settings.DOCUMENT_TEMPLATE_ROOT) / "pv_affectation" / template_name


def get_or_create_pv(dossier, administration):
    pv_key = build_pv_key(dossier)
    pv, _created = PvAffectation.objects.get_or_create(
        pv_key=pv_key,
        defaults={
            "source_import_id": dossier.import_id,
            "num_dossier": dossier.num_dossier,
            "numero_pv": dossier.numero_pv,
            "administration_source_nom": dossier.administration_beneficiaire or "",
            "administration": administration,
            "template_name": PV_TEMPLATE_NAME,
        },
    )
    return pv


def relative_generated_path(kind, pv_key, suffix):
    return str(Path("pv_affectation") / kind / f"{pv_key}.{suffix}")


def absolute_generated_path(relative_path):
    return Path(settings.GENERATED_DOCUMENT_ROOT) / relative_path


def ensure_parent(path):
    path.parent.mkdir(parents=True, exist_ok=True)


def calculate_sha256(path):
    digest = hashlib.sha256()
    with open(path, "rb") as file_obj:
        for chunk in iter(lambda: file_obj.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def electronic_signature_statement(administration):
    return f"{ELECTRONIC_SIGNATURE_PREFIX}{administration.nom}"


def add_electronic_signature_footer(docx_path, administration):
    """Add the beneficiary signature statement to the document footer once."""
    statement = electronic_signature_statement(administration)
    document = Document(docx_path)

    for section in document.sections:
        footer = section.footer
        paragraph = next(
            (item for item in footer.paragraphs if item.text == statement),
            None,
        )
        if paragraph is None:
            if len(footer.paragraphs) == 1 and not footer.paragraphs[0].text:
                paragraph = footer.paragraphs[0]
            else:
                paragraph = footer.add_paragraph()
            paragraph.add_run(statement)

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.italic = True

    document.save(docx_path)


def render_docx(dossier, administration, pv):
    template_path = get_template_path(pv.template_name)
    if not template_path.exists():
        raise FileNotFoundError(f"PV template not found: {template_path}")

    docx_relative_path = relative_generated_path("docx", pv.pv_key, "docx")
    docx_path = absolute_generated_path(docx_relative_path)
    ensure_parent(docx_path)

    document = DocxTemplate(template_path)
    document.render(build_pv_context(dossier, administration))
    document.save(docx_path)

    pv.generated_docx = docx_relative_path
    return docx_path


def convert_docx_to_pdf(docx_path, pv):
    pdf_relative_path = relative_generated_path("pdf", pv.pv_key, "pdf")
    pdf_path = absolute_generated_path(pdf_relative_path)
    ensure_parent(pdf_path)

    command = [
        settings.LIBREOFFICE_PATH,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(pdf_path.parent),
        str(docx_path),
    ]

    try:
        result = subprocess.run(
            command,
            check=False,
            capture_output=True,
            text=True,
            timeout=60,
        )
    except (FileNotFoundError, subprocess.TimeoutExpired) as exc:
        if settings.PV_ALLOW_DEVELOPMENT_PDF_FALLBACK:
            write_development_pdf_fallback(pdf_path, pv)
        else:
            raise ImproperlyConfigured(
                "LibreOffice headless is required for DOCX to PDF conversion."
            ) from exc
    else:
        if result.returncode != 0:
            if settings.PV_ALLOW_DEVELOPMENT_PDF_FALLBACK:
                write_development_pdf_fallback(pdf_path, pv)
            else:
                raise RuntimeError(result.stderr or result.stdout)

    if not pdf_path.exists():
        converted_path = docx_path.with_suffix(".pdf")
        if converted_path.exists():
            converted_path.replace(pdf_path)
        elif settings.PV_ALLOW_DEVELOPMENT_PDF_FALLBACK:
            write_development_pdf_fallback(pdf_path, pv)
        else:
            raise RuntimeError("LibreOffice did not produce the expected PDF file.")

    pv.generated_pdf = pdf_relative_path
    return pdf_path


def write_development_pdf_fallback(pdf_path, pv):
    text = f"PV d'affectation - {pv.num_dossier or ''} - {pv.numero_pv or ''}"
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"""BT
/F1 18 Tf
72 760 Td
({escaped}) Tj
ET
"""
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        f"<< /Length {len(stream.encode('latin-1'))} >>\nstream\n{stream}endstream",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    content = b"%PDF-1.4\n"
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(content))
        content += f"{index} 0 obj\n{obj}\nendobj\n".encode("latin-1")

    xref_offset = len(content)
    content += f"xref\n0 {len(objects) + 1}\n".encode("latin-1")
    content += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        content += f"{offset:010d} 00000 n \n".encode("latin-1")

    content += (
        f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    ).encode("latin-1")
    pdf_path.write_bytes(content)


def generate_pv_pdf(dossier, administration):
    pv = get_or_create_pv(dossier, administration)

    if pv.is_signed:
        return pv

    docx_path = render_docx(dossier, administration, pv)
    pdf_path = convert_docx_to_pdf(docx_path, pv)

    pv.pdf_hash_sha256 = calculate_sha256(pdf_path)
    pv.generated_at = timezone.now()
    pv.save(
        update_fields=[
            "generated_docx",
            "generated_pdf",
            "pdf_hash_sha256",
            "generated_at",
        ]
    )
    return pv


def generate_signed_pv_pdf(pv, administration):
    """Stamp the generated PV, regenerate its PDF, and return its final hash."""
    if not pv.generated_docx:
        raise RuntimeError("Le DOCX du PV doit etre genere avant sa signature.")

    docx_path = absolute_generated_path(pv.generated_docx)
    if not docx_path.exists():
        raise FileNotFoundError("Le DOCX genere du PV est introuvable.")

    add_electronic_signature_footer(docx_path, administration)
    pdf_path = convert_docx_to_pdf(docx_path, pv)
    return calculate_sha256(pdf_path)
