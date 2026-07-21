from django.contrib.auth import get_user_model
from django.contrib.auth.hashers import make_password
from django.core import mail
from django.db import connection
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from docx import Document
import tempfile
import shutil

from .models import (
    AdministrationBeneficiaire,
    OtpCode,
    PvAffectation,
    SignatureOtpPv,
    TableFaitAffectationDatalab,
)
from .services.pv_generation import absolute_generated_path, generate_pv_pdf


User = get_user_model()


class DossierListAccessTests(TestCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        with connection.schema_editor() as schema_editor:
            schema_editor.create_model(TableFaitAffectationDatalab)

    @classmethod
    def tearDownClass(cls):
        with connection.schema_editor() as schema_editor:
            schema_editor.delete_model(TableFaitAffectationDatalab)
        super().tearDownClass()

    def setUp(self):
        self.education = AdministrationBeneficiaire.objects.create(
            nom="Education Nationale"
        )
        self.sports = AdministrationBeneficiaire.objects.create(
            nom="Jeunesse Et Sports"
        )
        self.education_user = User.objects.create_user(
            username="education_user",
            password="StrongPass123!",
            role=User.ROLE_CONSULTATION,
            administration=self.education,
        )
        self.admin_dde = User.objects.create_user(
            username="admin_dde",
            password="StrongPass123!",
            role=User.ROLE_ADMIN_DDE,
            is_staff=True,
            is_superuser=True,
        )
        TableFaitAffectationDatalab.objects.create(
            num_dossier="EDU-001",
            administration_beneficiaire="Education Nationale",
            denomination_projet="Ecole primaire",
            statut_dossier="PV etabli",
            statut_pv="Valide",
        )
        TableFaitAffectationDatalab.objects.create(
            num_dossier="SPORT-001",
            administration_beneficiaire="Jeunesse Et Sports",
            denomination_projet="Complexe sportif",
            statut_dossier="En cours",
            statut_pv="",
        )

    def test_dossier_list_requires_login(self):
        response = self.client.get(reverse("affectations:dossier_list"))

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], "/login/?next=/dossiers/")

    def test_user_sees_only_dossiers_for_own_administration(self):
        self.client.force_login(self.education_user)

        response = self.client.get(reverse("affectations:dossier_list"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "EDU-001")
        self.assertContains(response, "Ecole primaire")
        self.assertNotContains(response, "SPORT-001")
        self.assertNotContains(response, "Complexe sportif")

    def test_admin_dde_cannot_access_extranet_dossier_list(self):
        self.client.force_login(self.admin_dde)

        response = self.client.get(reverse("affectations:dossier_list"))

        self.assertEqual(response.status_code, 403)


class PvAffectationFlowTests(TestCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        with connection.schema_editor() as schema_editor:
            schema_editor.create_model(TableFaitAffectationDatalab)

    @classmethod
    def tearDownClass(cls):
        with connection.schema_editor() as schema_editor:
            schema_editor.delete_model(TableFaitAffectationDatalab)
        super().tearDownClass()

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.addCleanup(lambda: shutil.rmtree(self.tmpdir, ignore_errors=True))
        self.template_root = f"{self.tmpdir}/document_templates"
        self.generated_root = f"{self.tmpdir}/generated_documents"
        self.override = override_settings(
            DOCUMENT_TEMPLATE_ROOT=self.template_root,
            GENERATED_DOCUMENT_ROOT=self.generated_root,
            PV_ALLOW_DEVELOPMENT_PDF_FALLBACK=True,
            PV_PRINT_OTP_TO_CONSOLE=False,
            EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend",
            DEFAULT_FROM_EMAIL="pv-otp@example.test",
        )
        self.override.enable()
        self.addCleanup(self.override.disable)
        self.create_template()

        self.education = AdministrationBeneficiaire.objects.create(
            nom="Education Nationale"
        )
        self.sports = AdministrationBeneficiaire.objects.create(
            nom="Jeunesse Et Sports"
        )
        self.consultation_user = User.objects.create_user(
            username="consultation",
            password="StrongPass123!",
            role=User.ROLE_CONSULTATION,
            administration=self.education,
        )
        self.signataire = User.objects.create_user(
            username="signataire",
            password="StrongPass123!",
            email="signataire@example.test",
            role=User.ROLE_SIGNATAIRE,
            administration=self.education,
            peut_signer=True,
        )
        self.other_signataire = User.objects.create_user(
            username="other_signataire",
            password="StrongPass123!",
            role=User.ROLE_SIGNATAIRE,
            administration=self.sports,
            peut_signer=True,
        )
        self.admin_dde = User.objects.create_user(
            username="pv_admin_dde",
            password="StrongPass123!",
            role=User.ROLE_ADMIN_DDE,
            is_staff=True,
            is_superuser=True,
        )

        self.ready_dossier = TableFaitAffectationDatalab.objects.create(
            num_dossier="EDU-PV-001",
            administration_beneficiaire="Education Nationale",
            denomination_projet="Ecole primaire",
            numero_pv="PV-001",
            type_pv="PV d'affectation",
            statut_pv="Signé par DR",
        )
        self.ready_duplicate = TableFaitAffectationDatalab.objects.create(
            num_dossier="EDU-PV-001",
            administration_beneficiaire="Education Nationale",
            denomination_projet="Ecole primaire - ligne duplicate",
            numero_pv="PV-001",
            type_pv="PV d'affectation",
            statut_pv="Signé par DR",
        )
        self.not_ready_dossier = TableFaitAffectationDatalab.objects.create(
            num_dossier="EDU-PV-002",
            administration_beneficiaire="Education Nationale",
            denomination_projet="College",
            numero_pv="PV-002",
            statut_pv="Validé",
        )

    def create_template(self):
        from pathlib import Path

        path = Path(self.template_root) / "pv_affectation" / "pv_affectation_template.docx"
        path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_paragraph("PV {{ num_dossier }} {{ administration_nom_fr }}")
        document.save(path)

    def test_pv_buttons_hidden_when_status_not_ready(self):
        self.client.force_login(self.signataire)

        response = self.client.get(
            reverse("affectations:dossier_detail", args=[self.not_ready_dossier.import_id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "Consulter le PV")
        self.assertNotContains(response, "Signer par OTP")

    def test_pv_buttons_hidden_for_consultation_user(self):
        self.client.force_login(self.consultation_user)

        response = self.client.get(
            reverse("affectations:dossier_detail", args=[self.ready_dossier.import_id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "Consulter le PV")
        self.assertNotContains(response, "Signer par OTP")

    def test_pv_buttons_shown_for_authorized_signataire(self):
        self.client.force_login(self.signataire)

        response = self.client.get(
            reverse("affectations:dossier_detail", args=[self.ready_dossier.import_id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Consulter le PV")
        self.assertContains(response, "Signer par OTP")

    def test_other_administration_cannot_access_dossier_or_pv(self):
        self.client.force_login(self.other_signataire)

        detail_response = self.client.get(
            reverse("affectations:dossier_detail", args=[self.ready_dossier.import_id])
        )
        pdf_response = self.client.get(
            reverse("affectations:pv_pdf", args=[self.ready_dossier.import_id])
        )

        self.assertEqual(detail_response.status_code, 403)
        self.assertEqual(pdf_response.status_code, 403)

    def test_admin_dde_cannot_access_extranet_pv_views(self):
        self.client.force_login(self.admin_dde)

        response = self.client.get(
            reverse("affectations:pv_pdf", args=[self.ready_dossier.import_id])
        )

        self.assertEqual(response.status_code, 403)

    def test_pdf_view_is_protected_and_does_not_expose_file_path(self):
        self.client.force_login(self.signataire)

        response = self.client.get(
            reverse("affectations:pv_pdf", args=[self.ready_dossier.import_id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "application/pdf")
        self.assertNotIn(self.generated_root, response["Content-Disposition"])

    def test_grouped_dossier_rows_share_one_pv_record(self):
        generate_pv_pdf(self.ready_dossier, self.education)
        generate_pv_pdf(self.ready_duplicate, self.education)

        self.assertEqual(PvAffectation.objects.count(), 1)

    def test_otp_request_creates_hashed_code_only(self):
        self.client.force_login(self.signataire)

        response = self.client.post(
            reverse("affectations:pv_otp_request", args=[self.ready_dossier.import_id])
        )
        otp = OtpCode.objects.get()

        self.assertEqual(response.status_code, 302)
        self.assertNotRegex(otp.code_hash, r"^\d{6}$")
        self.assertEqual(otp.attempts, 0)

    def test_otp_request_sends_email_only_to_the_signataire(self):
        self.client.force_login(self.signataire)

        response = self.client.post(
            reverse("affectations:pv_otp_request", args=[self.ready_dossier.import_id])
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(len(mail.outbox), 1)
        self.assertEqual(mail.outbox[0].to, [self.signataire.email])
        self.assertIn("code de verification", mail.outbox[0].subject)

    def test_new_otp_request_invalidates_the_previous_code(self):
        self.client.force_login(self.signataire)

        with override_settings(PV_OTP_REQUEST_COOLDOWN_SECONDS=0):
            self.client.post(
                reverse("affectations:pv_otp_request", args=[self.ready_dossier.import_id])
            )
            self.client.post(
                reverse("affectations:pv_otp_request", args=[self.ready_dossier.import_id])
            )

        first_otp, second_otp = OtpCode.objects.order_by("created_at")
        self.assertIsNotNone(first_otp.invalidated_at)
        self.assertEqual(first_otp.invalidation_reason, "replaced_by_new_request")
        self.assertIsNone(second_otp.invalidated_at)

    def test_wrong_otp_increments_attempts(self):
        pv = generate_pv_pdf(self.ready_dossier, self.education)
        OtpCode.objects.create(
            user=self.signataire,
            pv=pv,
            code_hash=make_password("123456"),
            expires_at=timezone.now() + timezone.timedelta(minutes=10),
        )
        self.client.force_login(self.signataire)

        response = self.client.post(
            reverse("affectations:pv_otp_verify", args=[self.ready_dossier.import_id]),
            {"otp_code": "000000"},
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(OtpCode.objects.get().attempts, 1)
        pv.refresh_from_db()
        self.assertFalse(pv.is_signed)

    def test_expired_otp_fails(self):
        pv = generate_pv_pdf(self.ready_dossier, self.education)
        OtpCode.objects.create(
            user=self.signataire,
            pv=pv,
            code_hash=make_password("123456"),
            expires_at=timezone.now() - timezone.timedelta(minutes=1),
        )
        self.client.force_login(self.signataire)

        response = self.client.post(
            reverse("affectations:pv_otp_verify", args=[self.ready_dossier.import_id]),
            {"otp_code": "123456"},
        )

        self.assertEqual(response.status_code, 302)
        pv.refresh_from_db()
        self.assertFalse(pv.is_signed)

    def test_used_otp_fails(self):
        pv = generate_pv_pdf(self.ready_dossier, self.education)
        OtpCode.objects.create(
            user=self.signataire,
            pv=pv,
            code_hash=make_password("123456"),
            expires_at=timezone.now() + timezone.timedelta(minutes=10),
            used_at=timezone.now(),
        )
        self.client.force_login(self.signataire)

        response = self.client.post(
            reverse("affectations:pv_otp_verify", args=[self.ready_dossier.import_id]),
            {"otp_code": "123456"},
        )

        self.assertEqual(response.status_code, 302)
        pv.refresh_from_db()
        self.assertFalse(pv.is_signed)

    def test_successful_otp_signature_marks_pv_signed_and_blocks_pdf(self):
        pv = generate_pv_pdf(self.ready_dossier, self.education)
        OtpCode.objects.create(
            user=self.signataire,
            pv=pv,
            code_hash=make_password("123456"),
            expires_at=timezone.now() + timezone.timedelta(minutes=10),
        )
        self.client.force_login(self.signataire)

        response = self.client.post(
            reverse("affectations:pv_otp_verify", args=[self.ready_dossier.import_id]),
            {"otp_code": "123456"},
            HTTP_USER_AGENT="Test agent",
        )
        pv.refresh_from_db()
        pdf_response = self.client.get(
            reverse("affectations:pv_pdf", args=[self.ready_dossier.import_id])
        )

        self.assertEqual(response.status_code, 302)
        self.assertTrue(pv.is_signed)
        self.assertIsNotNone(pv.signed_at)
        self.assertEqual(SignatureOtpPv.objects.count(), 1)
        self.assertEqual(SignatureOtpPv.objects.get().user_agent, "Test agent")
        signed_document = Document(absolute_generated_path(pv.generated_docx))
        footer_text = "\n".join(
            paragraph.text
            for section in signed_document.sections
            for paragraph in section.footer.paragraphs
        )
        self.assertIn(
            "Signé électroniquement par Education Nationale",
            footer_text,
        )
        self.assertEqual(pdf_response.status_code, 403)

        self.client.force_login(self.admin_dde)
        dde_pdf_response = self.client.get(
            reverse("affectations:dde_signed_pv_pdf", args=[pv.id])
        )
        self.assertEqual(dde_pdf_response.status_code, 200)
        self.assertEqual(dde_pdf_response["Content-Type"], "application/pdf")

        self.client.force_login(self.signataire)
        beneficiary_pdf_response = self.client.get(
            reverse("affectations:dde_signed_pv_pdf", args=[pv.id])
        )
        self.assertEqual(beneficiary_pdf_response.status_code, 403)
